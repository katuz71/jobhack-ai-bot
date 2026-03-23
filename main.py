import os, re, io, asyncio, logging, httpx
from PyPDF2 import PdfReader
from docx import Document
from dotenv import load_dotenv
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import CommandStart
from aiogram.types import InlineKeyboardButton, LinkPreviewOptions, BufferedInputFile
from aiogram.utils.keyboard import InlineKeyboardBuilder
from supabase import create_client, Client
from openai import AsyncOpenAI

# 1. НАСТРОЙКА ЛОГИРОВАНИЯ
load_dotenv(override=True)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

BOT_TOKEN = os.getenv("BOT_TOKEN", "").strip()
OPENAI_KEY = os.getenv("AI_TOKEN", "").strip()
S_URL = os.getenv("SUPABASE_URL", "").strip()
S_KEY = os.getenv("SUPABASE_KEY", "").strip()

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()
supabase: Client = create_client(S_URL, S_KEY)
ai_client = AsyncOpenAI(api_key=OPENAI_KEY, timeout=120.0)

MODEL = "gpt-4o-mini"
SEARCH_PROMPT = "Выдели 1-2 слова профессии для поиска на HH.ru. Только слова."
COVER_PROMPT = (
    "Ты — эксперт по найму. Напиши пробивное сопроводительное письмо. "
    "Используй структуру STAR, но НЕЯВНО (запрещено писать слова Situation, Task, Action, Result и делать заголовки для них). "
    "Текст должен быть связным, живым и уверенным. Фокус на цифрах и пользе для компании. "
    "ЗАПРЕЩЕНО использовать Markdown-заголовки (символы #). Только абзацы и списки (если нужно)."
)

# 2. ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
async def get_smart_query(text: str):
    try:
        res = await ai_client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "system", "content": SEARCH_PROMPT}, {"role": "user", "content": text}],
            max_tokens=15
        )
        return res.choices[0].message.content.strip().replace('"', '').replace('.', '')
    except Exception as e:
        logger.error(f"GPT Search Error: {e}")
        return "Маркетолог"

async def fetch_hh(query: str, page: int = 0):
    async with httpx.AsyncClient(timeout=40.0) as client:
        headers = {"User-Agent": "JobHackAI/1.0"}
        params = {"text": query, "per_page": 3, "page": page, "search_field": "name", "order_by": "relevance"}
        res = await client.get("https://api.hh.ru/vacancies", params=params, headers=headers)
        return res.json().get('items', [])

def create_docx(text: str, title: str):
    doc = Document()
    doc.add_heading('Сопроводительное письмо', 0)
    doc.add_heading(f'Вакансия: {title}', 1)
    doc.add_paragraph(text.replace('*', '')) 
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# 3. ОБРАБОТЧИКИ
@dp.message(CommandStart())
async def cmd_start(message: types.Message):
    try:
        supabase.table("users").upsert({"tg_id": message.from_user.id, "username": message.from_user.username}).execute()
    except: pass
    await message.answer("🚀 **JobHack AI готов к работе!**\n\nПришли свое PDF-резюме или напиши текстом: какой у тебя опыт?", parse_mode="Markdown")

@dp.message(F.document)
async def handle_pdf(message: types.Message):
    if not message.document.file_name.lower().endswith('.pdf'):
        return await message.answer("❌ Пожалуйста, пришли файл в формате PDF.")
    status = await message.answer("⏳ Читаю твое резюме...")
    try:
        file = await bot.get_file(message.document.file_id)
        file_bytes = io.BytesIO()
        await bot.download_file(file.file_path, destination=file_bytes)
        pdf = PdfReader(file_bytes)
        text = "\n".join([p.extract_text() for p in pdf.pages])
        supabase.table("users").update({"resume_text": text}).eq("tg_id", message.from_user.id).execute()
        q = await get_smart_query(text)
        await status.edit_text(f"✅ Резюме сохранено!\n🔎 Ищу лучшие вакансии по запросу: **{q}**", parse_mode="Markdown")
        await send_vacancies_block(message, q, page=0, is_edit=False, status_msg=status)
    except Exception as e:
        await status.edit_text("❌ Ошибка при чтении файла.")

@dp.message(F.text & ~F.text.startswith('/') & ~F.text.contains("hh.ru"))
async def handle_text(message: types.Message):
    try:
        supabase.table("users").upsert({"tg_id": message.from_user.id, "resume_text": message.text}).execute()
        status = await message.answer("🧠 Анализирую твой опыт...")
        q = await get_smart_query(message.text)
        await status.edit_text(f"🔎 Ищу подходящие вакансии: **{q}**", parse_mode="Markdown")
        await send_vacancies_block(message, q, page=0, is_edit=False, status_msg=status)
    except Exception as e:
        logger.error(f"Text error: {e}")

async def send_vacancies_block(message_or_call, query: str, page: int, is_edit: bool = False, status_msg: types.Message = None):
    vacs = await fetch_hh(query, page)
    if not vacs: 
        text = "😢 Больше подходящих вакансий по этому запросу не найдено."
        if is_edit:
            return await message_or_call.message.edit_text(text)
        else:
            return await (status_msg or message_or_call).edit_text(text) if status_msg else message_or_call.answer(text)
    
    builder = InlineKeyboardBuilder()
    text = f"🎯 **Топ-3 вакансии (Страница {page+1}):**\n\n"
    
    for i, v in enumerate(vacs):
        v_url = f"https://hh.ru/vacancy/{v['id']}"
        text += f"{i+1}. **[{v['name']}]({v_url})** в {v.get('employer', {}).get('name')}\n"
        builder.add(InlineKeyboardButton(text=f"Сгенерировать отклик {i+1}", callback_data=f"apply_{v['id']}"))
    
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="➡️ Показать еще вакансии", callback_data=f"more_{page+1}"))
    
    if is_edit:
        await message_or_call.message.edit_text(text, reply_markup=builder.as_markup(), parse_mode="Markdown", link_preview_options=LinkPreviewOptions(is_disabled=True))
    else:
        target = status_msg if status_msg else message_or_call
        if hasattr(target, 'edit_text'):
            await target.edit_text(text, reply_markup=builder.as_markup(), parse_mode="Markdown", link_preview_options=LinkPreviewOptions(is_disabled=True))
        else:
            await target.answer(text, reply_markup=builder.as_markup(), parse_mode="Markdown", link_preview_options=LinkPreviewOptions(is_disabled=True))

@dp.callback_query(F.data.startswith("more_"))
async def handle_more_vacancies(callback: types.CallbackQuery):
    await callback.answer()
    next_page = int(callback.data.split("_")[1])
    await callback.message.edit_text("⏳ Подгружаю следующую страницу вакансий...")
    try:
        res = supabase.table("users").select("resume_text").eq("tg_id", callback.from_user.id).execute()
        resume = res.data[0]['resume_text'] if res.data else "Кандидат"
        q = await get_smart_query(resume)
        await send_vacancies_block(callback, q, page=next_page, is_edit=True)
    except Exception as e:
        logger.error(f"More vacs error: {e}")
        await callback.message.edit_text("❌ Ошибка при загрузке новых вакансий.")

@dp.callback_query(F.data.startswith("apply_"))
async def handle_apply(callback: types.CallbackQuery):
    await callback.answer() 
    v_id = callback.data.split("_")[1]
    # ДОБАВЛЕН parse_mode="Markdown" сюда
    status_msg = await callback.message.answer(
        "⏳ **JobHack AI** анализирует вакансию и пишет оффер...\nОбычно это занимает 5-10 секунд.", 
        parse_mode="Markdown"
    )
    await generate_and_send_cover(callback.from_user.id, v_id, status_msg)

@dp.callback_query(F.data.startswith("reapply_"))
async def handle_reapply(callback: types.CallbackQuery):
    await callback.answer()
    v_id = callback.data.split("_")[1]
    # ДОБАВЛЕН parse_mode="Markdown" сюда
    await callback.message.edit_text(
        "⏳ **JobHack AI** придумывает новый вариант...\nПодожди пару секунд.", 
        parse_mode="Markdown"
    )
    await generate_and_send_cover(callback.from_user.id, v_id, callback.message)

async def generate_and_send_cover(user_id: int, v_id: str, message_to_edit: types.Message):
    v_url = f"https://hh.ru/vacancy/{v_id}"
    try:
        res = supabase.table("users").select("resume_text").eq("tg_id", user_id).execute()
        resume = res.data[0]['resume_text'] if res.data else "Кандидат"
        
        async with httpx.AsyncClient(timeout=40.0) as client:
            v = (await client.get(f"https://api.hh.ru/vacancies/{v_id}")).json()
        
        v_title, v_desc = v.get('name'), re.sub(r'<[^>]+>', ' ', v.get('description', ''))[:1500]
        
        ai_res = await ai_client.chat.completions.create(
            model=MODEL,
            temperature=0.7, 
            messages=[{"role": "system", "content": COVER_PROMPT}, 
                      {"role": "user", "content": f"Вак: {v_title}\nОписание: {v_desc}\nОпыт: {resume}"}]
        )
        
        raw_text = ai_res.choices[0].message.content
        
        final_response = (
            f"✅ **Твой отклик на {v_title}:**\n\n"
            f"{raw_text}\n\n"
            f"🔗 **Откликнуться тут:** [Перейти на HH.ru]({v_url})"
        )
        
        builder = InlineKeyboardBuilder()
        builder.add(InlineKeyboardButton(text="🔄 Переписать письмо", callback_data=f"reapply_{v_id}"))
        
        await message_to_edit.edit_text(
            final_response, 
            reply_markup=builder.as_markup(),
            parse_mode="Markdown",
            link_preview_options=LinkPreviewOptions(is_disabled=True)
        )
        
        doc_io = create_docx(raw_text, v_title)
        doc_file = BufferedInputFile(doc_io.read(), filename=f"JobHack_{v_id}.docx")
        await message_to_edit.answer_document(
            document=doc_file, 
            caption="📄 **Готовый документ!** Скачай, впиши свои контакты и отправляй.",
            parse_mode="Markdown"
        )
        
    except Exception as e:
        logger.error(f"❌ Ошибка генерации: {e}")
        await message_to_edit.edit_text("❌ Произошла ошибка генерации. Попробуй еще раз.")

# 4. ЗАПУСК
async def main():
    logger.info("🚀 Бот запущен (parse_mode добавлен во все статусы)")
    await dp.start_polling(bot)

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        logger.info("👋 Бот остановлен.")